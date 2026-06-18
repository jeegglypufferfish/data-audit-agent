"""Data Audit Agent verification engine.

Usage:
    python verify.py --report <report-file> --source <source-file> [<source-file> ...]

This script extracts numeric evidence from a report and one or more source files,
compares them with the rules defined in the team-agent docs, and writes an Excel
file that follows the verification template column layout.
"""

from __future__ import annotations

import argparse
import math
import re
import sys
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Sequence
from xml.etree import ElementTree as ET

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document
import pdfplumber


TEMPLATE_COLUMNS = [
    "항목",
    "보고서 수치",
    "단위",
    "보고서 위치",
    "보고서 수치 출처",
    "검증여부",
    "검증·비교 수치",
    "단위",
    "검증·비교 수치 위치",
    "검증·비교 수치 출처",
]

RESULT_LABELS = {
    "match": "검증완료(일치)",
    "mismatch": "검증완료(불일치)",
    "review": "확인 필요",
}

QUALITATIVE_EXCLUDES = {"낮음", "높음", "보통", "중간", "상", "하"}
APPROXIMATE_TERMS = ("약", "대략", "정도", "절반")
BIG_UNIT_TOKENS = ("조", "억", "만", "천", "백")
YEAR_LIKE_UNITS = {"년", "년도", "연도", "학년", "회계연도"}

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx"}
NUMERIC_PATTERN = re.compile(
    r"(?:[-+]?\d{1,3}(?:,\d{3})+|[-+]?\d+)(?:\.\d+)?",
)
YEAR_PATTERN = re.compile(r"[‘'’]?\s*(\d{2,4})\s*년")
APPROX_PATTERN = re.compile(r"절반")
GENERIC_LABELS = {
    "",
    "page",
    "pg",
    "p",
    "sheet",
    "시트",
    "table",
    "표",
    "본문",
    "문단",
    "행",
    "row",
    "col",
    "column",
}


@dataclass
class Block:
    location: str
    text: str
    reference: str = ""


@dataclass
class Candidate:
    id: str
    label: str
    period: str
    value: float
    unit: str
    decimals: int
    value_text: str
    location: str
    source_text: str
    evidence: str
    role: str
    approximate: bool = False


def normalize_text(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"\s+", " ", text)
    text = text.replace("|", " | ")
    text = text.replace("·", " ")
    text = text.replace("•", " ")
    text = text.replace("“", "").replace("”", "").replace('"', "").replace("'", "")
    return text.strip()


def normalize_key(value: object) -> str:
    text = normalize_text(value).lower()
    text = re.sub(r"[^\w가-힣%.\-/\s]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_qualitative_only(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return True
    return normalized in QUALITATIVE_EXCLUDES


def has_approximate_cue(text: str) -> bool:
    normalized = normalize_text(text)
    return any(term in normalized for term in APPROXIMATE_TERMS) or bool(APPROX_PATTERN.search(normalized))


def infer_unit(token: str) -> str:
    text = normalize_text(token).replace(" ", "").lower()
    text = re.sub(r"^[|,/:;·\-\(\)\[\]\{\}]+", "", text)
    if not text:
        return ""

    for unit in ("%p", "%", "원", "건", "회", "개", "세", "점", "명", "달러", "usd", "p", "년", "월", "일", "분기", "호", "차"):
        if text.startswith(unit):
            return unit

    for scale in ("조", "억", "만", "천", "백"):
        if text.startswith(scale):
            remainder = text[len(scale):]
            for unit in ("명", "건", "회", "개", "원", "달러", "usd", "세", "점", "년", "월", "일"):
                if remainder.startswith(unit):
                    return scale + unit
            return scale

    return ""


def normalize_year(value: str) -> str:
    digits = re.sub(r"\D", "", value or "")
    if not digits:
        return ""
    if len(digits) == 2:
        return f"20{digits}"
    return digits


def display_file_name(name: str) -> str:
    return re.sub(r"^(?:report|source)_\d+_", "", name, flags=re.IGNORECASE)


def compact_source_text(text: str) -> str:
    normalized = normalize_text(text)
    normalized = re.sub(r"\s*\(\s*", "(", normalized)
    normalized = re.sub(r"\s*\)\s*", ")", normalized)
    normalized = re.sub(r"\s*:\s*", ":", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalize_source_label(normalized.strip(" -:;,.。"))


def normalize_source_label(text: str) -> str:
    normalized = normalize_text(text)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    if re.fullmatch(r"KOSIS\s*대표\s*홈페이지", normalized, flags=re.IGNORECASE):
        return "KOSIS 대표홈페이지"
    if re.fullmatch(r"KOSIS\s*대표홈페이지", normalized, flags=re.IGNORECASE):
        return "KOSIS 대표홈페이지"
    return normalized


def extract_report_source(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return "출처 기재 안되어있음"

    patterns = [
        r"KOSIS\s*대표\s*홈페이지",
        r"이지원\s*\(\s*5091\s*\)",
        r"정우진\s*2019\s*논문",
        r"출처\s*기재\s*안되어있음",
        r"KOSIS\s*내부파일",
        r"출처\s*:\s*어디서부터\s*왔을까",
        r"알\s*수\s*없음",
    ]

    contextual = re.search(
        r"([가-힣A-Za-z0-9][가-힣A-Za-z0-9\s·\-_()]{1,60}?)(?:에 따르면|에 의하면|에 의거하면|에 따른|에 따른다면)",
        normalized,
    )
    if contextual:
        candidate = compact_source_text(contextual.group(1))
        if candidate and normalize_key(candidate) not in GENERIC_LABELS:
            return candidate

    for pattern in patterns:
        match = re.search(pattern, normalized, flags=re.IGNORECASE)
        if match:
            return compact_source_text(match.group(0))

    explicit = re.search(r"출처\s*[:：]\s*([^|/]+)", normalized)
    if explicit:
        candidate = compact_source_text(explicit.group(1))
        if candidate:
            return candidate

    if "출처 기재 안되어있음" in normalized:
        return "출처 기재 안되어있음"

    return "출처 기재 안되어있음"


def extract_sentence_around(text: str, start_index: int, end_index: int) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return ""

    left = 0
    right = len(normalized)

    def is_sentence_break(index: int) -> bool:
        char = normalized[index]
        if char not in ".!?":
            return False
        if char == ".":
            prev_char = normalized[index - 1] if index > 0 else ""
            next_char = normalized[index + 1] if index + 1 < len(normalized) else ""
            if prev_char.isdigit() and next_char.isdigit():
                return False
        return True

    for index in range(max(0, start_index - 1), -1, -1):
        if is_sentence_break(index):
            left = index + 1
            break

    for index in range(min(len(normalized) - 1, end_index), len(normalized)):
        if is_sentence_break(index):
            right = index + 1
            break

    return normalized[left:right].strip()


def extract_report_source_for_candidate(text: str, start_index: int, end_index: int, fallback: str) -> str:
    sentence = extract_sentence_around(text, start_index, end_index)
    if sentence:
        return extract_report_source(sentence)

    return fallback or "출처 기재 안되어있음"


def extract_period_from_text(text: str) -> str:
    matches = list(YEAR_PATTERN.finditer(normalize_text(text)))
    if not matches:
        plain_years = re.findall(r"\b(19\d{2}|20\d{2})\b", normalize_text(text))
        return plain_years[-1] if plain_years else ""
    return normalize_year(matches[-1].group(1))


def format_numeric_text(value: float, decimals: int) -> str:
    rounded = round_to(value, decimals)
    return f"{int(rounded)}" if float(rounded).is_integer() else f"{rounded}"


def parse_numeric_token(line: str, match: re.Match[str]) -> Candidate | None:
    number_text = match.group(0)
    suffix = normalize_text(line[match.end(): match.end() + 24])
    unit = infer_unit(suffix)
    cleaned = number_text.replace(",", "")
    try:
        numeric = float(cleaned)
    except ValueError:
        return None
    decimals = len(cleaned.split(".")[1]) if "." in cleaned else 0
    return Candidate(
        id="",
        label="",
        period="",
        value=numeric,
        unit=unit,
        decimals=decimals,
        value_text=cleaned,
        location="",
        source_text="",
        evidence="",
        role="",
    )


def is_year_like_candidate(line: str, match: re.Match[str], token: Candidate) -> bool:
    raw_digits = re.sub(r"\D", "", match.group(0) or "")
    if not raw_digits:
        return False

    try:
        numeric_int = int(float(token.value))
    except Exception:
        numeric_int = None

    if token.unit in YEAR_LIKE_UNITS:
        return True

    if len(raw_digits) == 4 and numeric_int is not None and 1900 <= numeric_int <= 2099:
        return True

    if numeric_int is not None and 0 <= numeric_int <= 99:
        window_start = max(0, match.start() - 2)
        window_end = min(len(line), match.end() + 4)
        window = normalize_text(line[window_start:window_end])
        if re.search(r"(?:년|년도|연도|학년|회계연도)", window):
            return True

    if match.start() > 0:
        before = line[match.start() - 1]
        if before in "([" and len(raw_digits) == 4 and numeric_int is not None and 1900 <= numeric_int <= 2099:
            return True

    return False


def candidate_signature(candidate: Candidate) -> tuple[str, str, str, str, str]:
    return (
        candidate.role,
        normalize_key(candidate.label),
        re.sub(r"\s+", "", candidate.value_text or ""),
        normalize_key(candidate.unit),
        "1" if candidate.approximate else "0",
    )


DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def extract_xml_node_text(node: ET.Element) -> str:
    parts: list[str] = []
    for child in node.iter():
        tag = child.tag.split("}")[-1] if isinstance(child.tag, str) and "}" in child.tag else child.tag
        if tag == "t":
            parts.append(child.text or "")
        elif tag in {"tab", "br", "cr"}:
            parts.append(" ")
    return normalize_text("".join(parts))


def extract_docx_xml_blocks(xml_bytes: bytes, location_prefix: str) -> list[Block]:
    blocks: list[Block] = []
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return blocks

    paragraphs = root.findall(".//w:p", DOCX_NS)
    for paragraph_index, paragraph in enumerate(paragraphs, start=1):
        text = extract_xml_node_text(paragraph)
        if text:
            blocks.append(Block(location=f"{location_prefix} 문단 {paragraph_index}", text=text))

    tables = root.findall(".//w:tbl", DOCX_NS)
    for table_index, table in enumerate(tables, start=1):
        rows = table.findall(".//w:tr", DOCX_NS)
        for row_index, row in enumerate(rows, start=1):
            cells = []
            for cell in row.findall("./w:tc", DOCX_NS):
                cell_text = extract_xml_node_text(cell)
                if cell_text:
                    cells.append(cell_text)
            if cells:
                blocks.append(Block(location=f"{location_prefix} 표 {table_index} 행 {row_index}", text=" | ".join(cells)))

    return blocks


def unit_category(unit: str) -> str:
    normalized = normalize_key(unit)
    if not normalized:
        return "none"
    if any(item in normalized for item in ("%", "p", "%p", "퍼센트", "퍼센트포인트")):
        return "percent"
    if any(item in normalized for item in ("건", "명", "회", "개", "점", "세")):
        return "count"
    if any(item in normalized for item in ("조", "억", "만", "천", "백", "원", "usd", "달러", "$")):
        return "scale"
    return "other"


def unit_multiplier(unit: str) -> float:
    normalized = normalize_key(unit)
    if "조" in normalized:
        return 1e12
    if "억" in normalized:
        return 1e8
    if "만" in normalized:
        return 1e4
    if "천" in normalized:
        return 1e3
    if "백" in normalized:
        return 1e2
    return 1.0


def scale_token_from_text(text: str) -> str:
    normalized = normalize_key(text)
    for token in BIG_UNIT_TOKENS:
        if token in normalized:
            return token
    return ""


def pick_comparison_unit(report_candidate: Candidate, source_candidate: Candidate) -> str:
    report_unit = report_candidate.unit or ""
    source_unit = source_candidate.unit or ""
    report_category = unit_category(report_unit)
    source_category = unit_category(source_unit)

    if report_category == "percent" or source_category == "percent":
        return report_unit or source_unit or "%"

    if report_category == source_category == "count":
        return report_unit or source_unit or ""

    report_scale = scale_token_from_text(report_unit)
    source_scale = scale_token_from_text(source_unit)
    if report_scale:
        return report_scale
    if source_scale:
        return source_scale

    report_base = abs(report_candidate.value * unit_multiplier(report_unit))
    source_base = abs(source_candidate.value * unit_multiplier(source_unit))
    max_base = max(report_base, source_base)

    if max_base >= 1e12:
        return "조"
    if max_base >= 1e8:
        return "억"
    if max_base >= 1e4:
        return "만"
    if report_category == source_category and report_category != "none":
        return report_unit or source_unit or ""
    if report_category != "none":
        return report_unit
    if source_category != "none":
        return source_unit
    return ""


def round_to(value: float, decimals: int) -> float:
    factor = 10**decimals
    return round((value + sys.float_info.epsilon) * factor) / factor


def truncate_text(value: str, max_length: int = 140) -> str:
    text = normalize_text(value)
    if len(text) <= max_length:
        return text
    return text[: max_length - 1] + "…"


def derive_label(prefix: str, full_text: str) -> str:
    text = normalize_text(prefix) or normalize_text(full_text)
    if not text:
        return ""

    sentence_parts = re.split(r"[.!?。]\s*", text)
    text = sentence_parts[-1] if sentence_parts else text
    clause_parts = re.split(r"(?:그리고|이고|및|또는|그러나|하지만)\s*", text)
    text = clause_parts[-1] if clause_parts else text
    if "|" in text:
        pipe_parts = [piece.strip() for piece in text.split("|") if piece.strip()]
        if pipe_parts:
            text = pipe_parts[0]

    text = re.sub(r"^.*?(?:에 따르면|에 의하면|에 의거하면)\s*", "", text)

    year_matches = list(YEAR_PATTERN.finditer(text))
    if year_matches:
        text = text[year_matches[-1].end():]

    text = re.sub(r"^(?:KOSIS 대표 홈페이지에 따르면|우리 나라|총|대략|약|정도|약간)\s*", "", text)
    particle_match = re.search(r"(.+?)(?:은|는|이|가|을|를)\s*(?:총|대략|약|정도)?$", text)
    if particle_match:
        text = particle_match.group(1)
    text = re.sub(r"[\d,]+", "", text)
    text = re.sub(r"[|/:=,·•()]+", " ", text)
    text = re.sub(r"(?:은|는|이|가|을|를|의|이다|이야|입니다|다|고)\s*$", "", text)
    text = re.sub(r"\s+", "", text)
    if not text or normalize_key(text) in GENERIC_LABELS:
        return ""
    return text


def extract_blocks_from_plain_text(text: str, source_label: str) -> list[Block]:
    blocks: list[Block] = []
    for index, line in enumerate(str(text or "").splitlines(), start=1):
        normalized = normalize_text(line)
        if normalized:
            blocks.append(Block(location=f"{source_label} 본문 {index}", text=normalized))
    return blocks


def extract_pdf_blocks(file_path: Path) -> list[Block]:
    blocks: list[Block] = []
    seen: set[tuple[int, str]] = set()
    with pdfplumber.open(str(file_path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            for line_index, line in enumerate(text.splitlines(), start=1):
                normalized = normalize_text(line)
                if not normalized:
                    continue
                key = (page_index, normalized)
                if key in seen:
                    continue
                seen.add(key)
                blocks.append(Block(location=f"{file_path.name} pg {page_index} 본문 {line_index}", text=normalized))

            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for table_index, table in enumerate(tables, start=1):
                for row_index, row in enumerate(table, start=1):
                    cells = [normalize_text(cell) for cell in row if normalize_text(cell)]
                    if not cells:
                        continue
                    joined = " | ".join(cells)
                    key = (page_index, joined)
                    if key in seen:
                        continue
                    seen.add(key)
                    blocks.append(Block(location=f"{file_path.name} pg {page_index} 표 {table_index} 행 {row_index}", text=joined))
    return blocks


def extract_docx_blocks(file_path: Path) -> list[Block]:
    document = Document(str(file_path))
    blocks: list[Block] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)
        if text:
            blocks.append(Block(location="보고서 내 pg 1", text=text, reference=extract_report_source(text)))

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if not cells:
                continue
            row_text = " | ".join(cells)
            blocks.append(Block(location="보고서 내 pg 1", text=row_text, reference=extract_report_source(row_text)))

    try:
        with zipfile.ZipFile(str(file_path)) as archive:
            extra_names = [
                name
                for name in archive.namelist()
                if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
                or name in {"word/footnotes.xml", "word/endnotes.xml"}
            ]
            for extra_name in sorted(extra_names):
                extra_blocks = extract_docx_xml_blocks(
                    archive.read(extra_name),
                    f"{file_path.name} {Path(extra_name).stem}",
                )
                for block in extra_blocks:
                    blocks.append(block)
    except Exception:
        pass
    return blocks


def extract_legacy_doc_blocks(file_path: Path) -> list[Block]:
    try:
        return extract_blocks_from_plain_text(file_path.read_text(encoding="utf-8", errors="ignore"), file_path.name)
    except Exception:
        return []


def extract_xlsx_blocks(file_path: Path) -> list[Block]:
    workbook = load_workbook(str(file_path), data_only=True)
    blocks: list[Block] = []

    for sheet in workbook.worksheets:
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [normalize_text(value) for value in row if normalize_text(value)]
            if not values:
                continue
            row_text = " | ".join(values)
            source_reference = values[-1] if len(values) >= 4 else ""
            source_location = f"{display_file_name(file_path.name)} {sheet.title}!row {row_index}"
            blocks.append(
                Block(
                    location=source_location,
                    text=row_text,
                    reference=normalize_source_label(source_reference) if source_reference else "",
                )
            )
    return blocks


def extract_blocks_from_file(file_path: Path) -> list[Block]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_blocks(file_path)
    if ext == ".docx":
        return extract_docx_blocks(file_path)
    if ext == ".xlsx":
        return extract_xlsx_blocks(file_path)
    if ext == ".doc":
        return extract_legacy_doc_blocks(file_path)
    return extract_blocks_from_plain_text(file_path.read_text(encoding="utf-8", errors="ignore"), file_path.name)


def extract_numeric_candidates(blocks: Sequence[Block], role: str) -> list[Candidate]:
    candidates: list[Candidate] = []
    generic_labels = GENERIC_LABELS
    seen_signatures: set[tuple[str, str, str, str, str]] = set()

    for block_index, block in enumerate(blocks):
        line = normalize_text(block.text)
        if not line:
            continue

        matches = list(NUMERIC_PATTERN.finditer(line))
        if not matches:
            approx = APPROX_PATTERN.search(line)
            if not approx:
                continue
            label = derive_label(line[: approx.start()], line)
            if not label or normalize_key(label) in generic_labels:
                continue
            candidate = Candidate(
                id=f"{role}-{block_index}-approx",
                label=label,
                period=extract_period_from_text(line),
                value=50.0,
                unit="%",
                decimals=0,
                value_text=normalize_text(line),
                location=block.location,
                source_text=line,
                evidence=block.reference,
                role=role,
                approximate=True,
            )
            signature = candidate_signature(candidate)
            if signature in seen_signatures:
                continue
            seen_signatures.add(signature)
            candidates.append(candidate)
            continue

        for match_index, match in enumerate(matches):
            token = parse_numeric_token(line, match)
            if not token:
                continue
            if is_year_like_candidate(line, match, token):
                continue
            label = derive_label(line[: match.start()], line)
            normalized_label = normalize_key(label)
            if not label or normalized_label in generic_labels:
                continue
            if normalized_label in QUALITATIVE_EXCLUDES or normalize_text(label) in QUALITATIVE_EXCLUDES:
                continue
            if len(normalized_label) < 2:
                continue
            nearby_text = normalize_text(line[max(0, match.start() - 18): min(len(line), match.end() + 18)])
            candidate = Candidate(
                id=f"{role}-{block_index}-{match_index}",
                label=label,
                period=extract_period_from_text(line[: match.start()]),
                value=token.value,
                unit=token.unit,
                decimals=token.decimals,
                value_text=token.value_text,
                location=block.location,
                source_text=line,
                evidence=extract_report_source_for_candidate(line, match.start(), match.end(), block.reference) if role == "report" else block.reference,
                role=role,
                approximate=has_approximate_cue(nearby_text),
            )
            signature = candidate_signature(candidate)
            if signature in seen_signatures:
                continue
            seen_signatures.add(signature)
            candidates.append(candidate)
    return candidates


def format_display_value(candidate: Candidate, display_unit: str) -> str:
    base_value = candidate.value * unit_multiplier(candidate.unit)
    numeric = base_value / unit_multiplier(display_unit)
    rounded = round_to(numeric, candidate.decimals)
    text = f"{int(rounded)}" if float(rounded).is_integer() else f"{rounded}"
    return f"{text}{display_unit}" if display_unit else text


def format_item_label(candidate: Candidate, fallback_period: str = "") -> str:
    label = re.sub(r"\s+", "", candidate.label or "")
    period = candidate.period or fallback_period
    if period and re.fullmatch(r"(?:19\d{2}|20\d{2}|\d{2})", period):
        period = ""
    return f"{label}({period})" if period else label


def compare_candidates(report_candidate: Candidate, source_candidate: Candidate) -> dict[str, str | bool]:
    report_category = unit_category(report_candidate.unit)
    source_category = unit_category(source_candidate.unit)
    display_unit = pick_comparison_unit(report_candidate, source_candidate)
    display_category = unit_category(display_unit)

    if report_category != "none" and source_category != "none" and report_category != source_category:
        return {
            "result": "검증 불가",
            "displayResult": RESULT_LABELS["review"],
            "badge": "unavailable",
            "note": "단위 범주가 달라 직접 비교할 수 없습니다.",
            "delta": "단위 범주 다름",
            "report_display": format_display_value(report_candidate, display_unit),
            "source_display": format_display_value(source_candidate, display_unit),
            "equal": False,
        }

    report_base = report_candidate.value * unit_multiplier(report_candidate.unit)
    source_base = source_candidate.value * unit_multiplier(source_candidate.unit)
    report_display_number = report_base / unit_multiplier(display_unit)
    source_display_number = source_base / unit_multiplier(display_unit)
    decimals = report_candidate.decimals
    report_rounded = round_to(report_display_number, decimals)
    source_rounded = round_to(source_display_number, decimals)
    exact_equal = report_rounded == source_rounded
    diff = source_rounded - report_rounded

    if display_unit == "%":
        diff_text = f"{diff:+g}%p"
    else:
        diff_text = f"{diff:+g}{display_unit}"

    approx_cue = (
        report_candidate.approximate
        or source_candidate.approximate
    )
    base_tolerance = 5 if display_category == "percent" else max(abs(source_rounded) * 0.05, 1)
    approx_equal = approx_cue and abs(diff) <= base_tolerance
    equal = exact_equal or approx_equal

    if equal:
        if exact_equal:
            note = "보고서 값과 Source 값이 반올림 기준상 일치합니다."
        else:
            note = "보고서의 근사 표현 기준상 일치합니다."
        result = "일치"
        badge = "match"
        delta = "-"
    else:
        note = f"보고서 값과 Source 값이 다릅니다. 차이 {diff_text}"
        result = "불일치"
        badge = "mismatch"
        delta = diff_text

    return {
        "result": result,
        "displayResult": RESULT_LABELS["match"] if result == "일치" else RESULT_LABELS["mismatch"],
        "badge": badge,
        "note": note,
        "delta": delta,
        "report_display": f"{int(report_rounded)}" if float(report_rounded).is_integer() else f"{report_rounded}",
        "source_display": f"{int(source_rounded)}" if float(source_rounded).is_integer() else f"{source_rounded}",
        "equal": equal,
    }


def similarity(a: str, b: str) -> float:
    set_a = set(normalize_key(a).split())
    set_b = set(normalize_key(b).split())
    if not set_a or not set_b:
        return 0.0
    overlap = len(set_a & set_b)
    return (2 * overlap) / (len(set_a) + len(set_b))


def match_validation_rows(report_candidates: Sequence[Candidate], source_candidates: Sequence[Candidate]) -> list[dict[str, str]]:
    used_source: set[int] = set()
    rows: list[dict[str, str]] = []

    for report_candidate in report_candidates:
        best_index = -1
        best_score = 0.0

        for index, source_candidate in enumerate(source_candidates):
            if index in used_source:
                continue
            label_score = similarity(report_candidate.label, source_candidate.label)
            text_score = similarity(report_candidate.source_text, source_candidate.source_text) * 0.4
            unit_bonus = 0.2 if report_candidate.unit and source_candidate.unit and normalize_key(report_candidate.unit) == normalize_key(source_candidate.unit) else 0.0
            score = (label_score * 0.6) + text_score + unit_bonus
            if score > best_score:
                best_score = score
                best_index = index

        if best_index < 0 or best_score < 0.3:
            result = "검증 불가" if report_candidate.approximate else "누락"
            rows.append(
                {
                    "item": format_item_label(report_candidate),
                    "reportValue": report_candidate.value_text,
                    "reportUnit": report_candidate.unit or "-",
                    "reportLocation": report_candidate.location,
                    "reportSource": report_candidate.evidence or "출처 기재 안되어있음",
                    "sourceValue": "-",
                    "sourceUnit": "-",
                    "sourceLocation": "-",
                    "sourceSource": "-",
                    "result": result,
                    "displayResult": RESULT_LABELS["review"],
                    "note": "근사 표현이 포함되어 있어 직접 비교가 어렵습니다." if report_candidate.approximate else "Source File에서 대응 항목을 찾지 못했습니다.",
                    "badge": "unavailable" if report_candidate.approximate else "missing",
                    "delta": "문맥 비교 필요" if report_candidate.approximate else "비교값 없음",
                }
            )
            continue

        used_source.add(best_index)
        source_candidate = source_candidates[best_index]
        comparison = compare_candidates(report_candidate, source_candidate)
        fallback_period = source_candidate.period or report_candidate.period
        rows.append(
            {
                "item": format_item_label(report_candidate, fallback_period),
                "reportValue": report_candidate.value_text,
                "reportUnit": report_candidate.unit or "-",
                "reportLocation": report_candidate.location,
                "reportSource": report_candidate.evidence or "출처 기재 안되어있음",
                "sourceValue": source_candidate.value_text,
                "sourceUnit": source_candidate.unit or "-",
                "sourceLocation": source_candidate.location,
                "sourceSource": source_candidate.evidence or "-",
                "result": comparison["result"],
                "displayResult": comparison["displayResult"],
                "note": f"{comparison['note']} / 보고서 위치: {report_candidate.location} / Source 위치: {source_candidate.location}",
                "badge": comparison["badge"],
                "delta": comparison["delta"],
            }
        )

    return rows


def run_verification(report_paths: Sequence[Path], source_paths: Sequence[Path]) -> dict:
    report_blocks: list[Block] = []
    for report_path in report_paths:
        report_blocks.extend(extract_blocks_from_file(report_path))
    source_blocks: list[Block] = []
    for source_path in source_paths:
        source_blocks.extend(extract_blocks_from_file(source_path))

    report_candidates = extract_numeric_candidates(report_blocks, "report")
    source_candidates = extract_numeric_candidates(source_blocks, "source")

    if not report_candidates and not source_candidates:
        return {
            "summary": {
                "total": 0,
                "match": 0,
                "mismatch": 0,
                "missing": 0,
                "unavailable": 0,
            },
            "rows": [],
            "report_count": len(report_candidates),
            "source_count": len(source_candidates),
            "notice": "보고서와 Source File에서 읽을 수 있는 수치형 데이터를 찾지 못했습니다.",
        }

    if report_candidates and not source_candidates:
        rows = [
            {
                "item": format_item_label(candidate),
                "reportValue": candidate.value_text,
                "reportUnit": candidate.unit or "-",
                "reportLocation": candidate.location,
                "reportSource": candidate.evidence or "출처 기재 안되어있음",
                "sourceValue": "-",
                "sourceUnit": "-",
                "sourceLocation": "-",
                "sourceSource": "-",
                "result": "검증 불가" if candidate.approximate else "누락",
                "displayResult": RESULT_LABELS["review"],
                "note": "Source File이 없어 직접 비교를 진행하지 못했습니다." if not candidate.approximate else "근사 표현이 포함되어 있어 직접 비교가 어렵습니다.",
                "badge": "unavailable" if candidate.approximate else "missing",
                "delta": "비교값 없음" if not candidate.approximate else "문맥 비교 필요",
            }
            for candidate in report_candidates
        ]
        summary = {
            "total": len(rows),
            "match": 0,
            "mismatch": 0,
            "missing": sum(1 for row in rows if row["result"] == "누락"),
            "unavailable": sum(1 for row in rows if row["result"] == "검증 불가"),
        }
        return {
            "summary": summary,
            "rows": rows,
            "report_count": len(report_candidates),
            "source_count": len(source_candidates),
            "notice": "Source File에서 대응값을 찾지 못했지만 보고서 수치 추출은 계속했습니다.",
        }

    rows = match_validation_rows(report_candidates, source_candidates)
    summary = {
        "total": len(rows),
        "match": sum(1 for row in rows if row["result"] == "일치"),
        "mismatch": sum(1 for row in rows if row["result"] == "불일치"),
        "missing": sum(1 for row in rows if row["result"] == "누락"),
        "unavailable": sum(1 for row in rows if row["result"] == "검증 불가"),
    }
    return {
        "summary": summary,
        "rows": rows,
        "report_count": len(report_candidates),
        "source_count": len(source_candidates),
        "notice": "",
    }


def ensure_output_path(output_dir: Path, base_name: str = "검증결과.xlsx") -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    candidate = output_dir / base_name
    if not candidate.exists():
        return candidate

    index = 1
    while True:
        numbered = output_dir / f"검증결과({index}).xlsx"
        if not numbered.exists():
            return numbered
        index += 1


def write_excel(result: dict, output_path: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "검증결과"

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    normal_font = Font(color="000000")
    red_font = Font(color="FF0000")
    center = Alignment(vertical="center", wrap_text=True)

    ws.append(TEMPLATE_COLUMNS)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center

    rows = result["rows"]
    for row in rows:
        excel_row = [
            row["item"],
            row["reportValue"],
            row["reportUnit"],
            row["reportLocation"],
            row["reportSource"],
            row["displayResult"],
            row["sourceValue"],
            row["sourceUnit"],
            row["sourceLocation"],
            row["sourceSource"],
        ]
        ws.append(excel_row)
        target = ws[ws.max_row]
        row_font = red_font if row["result"] == "불일치" else normal_font
        for cell in target:
            cell.font = row_font
            cell.alignment = center

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:J{ws.max_row}"

    width_map = {
        "A": 28,
        "B": 18,
        "C": 10,
        "D": 24,
        "E": 28,
        "F": 18,
        "G": 18,
        "H": 10,
        "I": 24,
        "J": 28,
    }
    for col, width in width_map.items():
        ws.column_dimensions[col].width = width

    if result["notice"]:
        ws.insert_rows(2)
        ws["A2"] = result["notice"]
        ws["A2"].font = Font(color="7A7A7A", italic=True)
        ws.merge_cells("A2:J2")

    wb.save(output_path)
    return output_path


def summarize_result(result: dict, output_path: Path) -> str:
    summary = result["summary"]
    return (
        f"검증 완료: 총 {summary['total']}건, "
        f"일치 {summary['match']}건, "
        f"불일치 {summary['mismatch']}건, "
        f"누락 {summary['missing']}건, "
        f"검증불가 {summary['unavailable']}건\n"
        f"출력 파일: {output_path}"
    )


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Data Audit Agent verification engine")
    parser.add_argument("--report", required=True, nargs="+", help="검증 대상 보고서 파일 경로")
    parser.add_argument("--source", required=True, nargs="+", help="원본 데이터 파일 경로(1개 이상)")
    parser.add_argument(
        "--output-dir",
        default=str(Path(__file__).resolve().parents[1] / "outputs"),
        help="검증결과.xlsx를 저장할 폴더 경로",
    )
    parser.add_argument("--output-file", default="", help="원하는 출력 파일명(비어 있으면 자동 번호)")
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    report_paths = [Path(item).expanduser().resolve() for item in args.report]
    source_paths = [Path(item).expanduser().resolve() for item in args.source]
    output_dir = Path(args.output_dir).expanduser().resolve()

    for report_path in report_paths:
        if not report_path.exists():
            print(f"[오류] 보고서 파일을 찾을 수 없습니다: {report_path}", file=sys.stderr)
            return 2
        if report_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            print(f"[오류] 지원하지 않는 보고서 파일 형식입니다: {report_path.name}", file=sys.stderr)
            return 2

    for source_path in source_paths:
        if not source_path.exists():
            print(f"[오류] 원본 데이터 파일을 찾을 수 없습니다: {source_path}", file=sys.stderr)
            return 2
        if source_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            print(f"[오류] 지원하지 않는 파일 형식입니다: {source_path.name}", file=sys.stderr)
            return 2

    result = run_verification(report_paths, source_paths)
    if args.output_file:
        output_path = output_dir / args.output_file
        if output_path.exists():
            output_path = ensure_output_path(output_dir, Path(args.output_file).name)
    else:
        output_path = ensure_output_path(output_dir)

    write_excel(result, output_path)
    print(summarize_result(result, output_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
